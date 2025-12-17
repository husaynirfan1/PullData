"""
DOCX parsing using python-docx.

Extracts text content and basic structure from Microsoft Word documents.
"""

import hashlib
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument

from pulldata.core.datatypes import Document, DocumentType
from pulldata.core.exceptions import (
    DocumentNotFoundError,
    ParsingError,
    UnsupportedFormatError,
)


class DOCXParser:
    """
    DOCX parser using python-docx.

    Extracts text content, paragraphs, and basic metadata from Word documents.
    """

    def __init__(self):
        """Initialize DOCX parser."""
        self.supported_extensions = {".docx"}

    def is_supported(self, file_path: str | Path) -> bool:
        """
        Check if file is a supported DOCX.

        Args:
            file_path: Path to file

        Returns:
            True if file is a DOCX, False otherwise
        """
        path = Path(file_path)
        return path.suffix.lower() in self.supported_extensions

    def parse(
        self,
        file_path: str | Path,
        extract_metadata: bool = True,
    ) -> tuple[Document, str]:
        """
        Parse a DOCX file.

        Args:
            file_path: Path to DOCX file
            extract_metadata: Whether to extract document metadata

        Returns:
            Tuple of (Document, full_text)
            - Document: Document metadata
            - full_text: Complete extracted text

        Raises:
            DocumentNotFoundError: If file doesn't exist
            UnsupportedFormatError: If file is not a DOCX
            ParsingError: If parsing fails
        """
        file_path = Path(file_path)

        # Validate file exists
        if not file_path.exists():
            raise DocumentNotFoundError(
                f"DOCX file not found: {file_path}",
                details={"path": str(file_path)},
            )

        # Validate file type
        if not self.is_supported(file_path):
            raise UnsupportedFormatError(
                f"File is not a DOCX: {file_path}",
                details={"path": str(file_path), "extension": file_path.suffix},
            )

        try:
            # Open DOCX
            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            # Combine all text
            full_text = "\n\n".join(paragraphs)

            # Hash content
            content_hash = hashlib.sha256(full_text.encode()).hexdigest()

            # Get file size
            file_size = file_path.stat().st_size

            # Extract metadata
            metadata = {}
            if extract_metadata:
                core_props = doc.core_properties
                if core_props:
                    # Extract available properties
                    if core_props.author:
                        metadata["author"] = core_props.author
                    if core_props.title:
                        metadata["title"] = core_props.title
                    if core_props.subject:
                        metadata["subject"] = core_props.subject
                    if core_props.created:
                        metadata["created"] = core_props.created.isoformat()
                    if core_props.modified:
                        metadata["modified"] = core_props.modified.isoformat()

            # Create Document
            document = Document(
                source_path=str(file_path),
                filename=file_path.name,
                doc_type=DocumentType.DOCX,
                content_hash=content_hash,
                file_size=file_size,
                num_pages=None,  # DOCX doesn't have fixed pages
                metadata=metadata,
            )

            return document, full_text

        except Exception as e:
            raise ParsingError(
                f"Failed to parse DOCX: {e}",
                details={"path": str(file_path), "error": str(e), "type": type(e).__name__},
            )

    def extract_paragraphs(self, file_path: str | Path) -> list[str]:
        """
        Extract paragraphs from DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of paragraph texts

        Raises:
            ParsingError: If extraction fails
        """
        file_path = Path(file_path)

        try:
            doc = DocxDocument(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            return paragraphs

        except Exception as e:
            raise ParsingError(
                f"Failed to extract paragraphs: {e}",
                details={"path": str(file_path), "error": str(e)},
            )

    def extract_tables(self, file_path: str | Path) -> list[list[list[str]]]:
        """
        Extract tables from DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of tables, where each table is a list of rows,
            and each row is a list of cell texts

        Raises:
            ParsingError: If extraction fails
        """
        file_path = Path(file_path)

        try:
            doc = DocxDocument(file_path)
            tables = []

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)

            return tables

        except Exception as e:
            raise ParsingError(
                f"Failed to extract tables: {e}",
                details={"path": str(file_path), "error": str(e)},
            )

    def count_paragraphs(self, file_path: str | Path) -> int:
        """
        Count non-empty paragraphs in DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            Number of paragraphs

        Raises:
            ParsingError: If counting fails
        """
        try:
            paragraphs = self.extract_paragraphs(file_path)
            return len(paragraphs)
        except Exception as e:
            raise ParsingError(
                f"Failed to count paragraphs: {e}",
                details={"path": str(file_path), "error": str(e)},
            )

    def has_tables(self, file_path: str | Path) -> bool:
        """
        Check if DOCX contains any tables.

        Args:
            file_path: Path to DOCX file

        Returns:
            True if document contains tables, False otherwise
        """
        try:
            doc = DocxDocument(file_path)
            return len(doc.tables) > 0
        except Exception:
            return False
